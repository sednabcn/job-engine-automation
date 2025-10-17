"""
File Readers Module
Handles reading various file formats (PDF, DOCX, TXT) for CV and job descriptions.
"""

import logging
import re
from pathlib import Path
from typing import Any, Dict

import docx
import PyPDF2

logger = logging.getLogger(__name__)


class FileReader:
    """Base class for file reading operations."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}

    @staticmethod
    def read_file(file_path: str) -> str:
        """
        Read file content based on file extension.

        Args:
        file_path: Path to the file

        Returns:
        Extracted text content

        Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file format is not supported
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in FileReader.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(FileReader.SUPPORTED_EXTENSIONS)}"
            )

        try:
            if extension == ".pdf":
                return PDFReader.read(file_path)
            elif extension in {".docx", ".doc"}:
                return DOCXReader.read(file_path)
            elif extension == ".txt":
                return TXTReader.read(file_path)
            else:
                # This should never happen due to earlier validation,
                # but ensures all paths return or raise
                raise ValueError(f"Unhandled file extension: {extension}")
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            raise

    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """
        Get file metadata.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with file information
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        stats = path.stat()

        return {
            "name": path.name,
            "extension": path.suffix,
            "size_bytes": stats.st_size,
            "size_kb": round(stats.st_size / 1024, 2),
            "modified": stats.st_mtime,
            "absolute_path": str(path.absolute()),
        }


class PDFReader:
    """Reader for PDF files."""

    @staticmethod
    def read(file_path: str) -> str:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Extracted text content
        """
        try:
            text_content = []

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {str(e)}")
                        continue

            full_text = "\n".join(text_content)

            # Clean up text
            full_text = PDFReader._clean_text(full_text)

            if not full_text.strip():
                logger.warning(f"No text extracted from PDF: {file_path}")

            return full_text

        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            raise ImportError(
                "PyPDF2 is required for PDF reading. " "Install with: pip install PyPDF2"
            )
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            raise

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean extracted PDF text."""
        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove page numbers pattern
        text = re.sub(r"\n\d+\n", "\n", text)
        # Fix common PDF extraction issues
        text = text.replace("\x00", "")
        return text.strip()


class DOCXReader:
    """Reader for DOCX/DOC files."""

    @staticmethod
    def read(file_path: str) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text content
        """
        try:
            text_content = []
            doc = docx.Document(file_path)

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_content.append(row_text)

            full_text = "\n".join(text_content)

            if not full_text.strip():
                logger.warning(f"No text extracted from DOCX: {file_path}")

            return full_text.strip()

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise ImportError(
                "python-docx is required for DOCX reading. " "Install with: pip install python-docx"
            )
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            raise


class TXTReader:
    """Reader for plain text files."""

    @staticmethod
    def read(file_path: str, encoding: str = "utf-8") -> str:
        """
        Read plain text file.

        Args:
            file_path: Path to text file
            encoding: File encoding (default: utf-8)

        Returns:
            File content
        """
        encodings_to_try = [encoding, "utf-8", "latin-1", "cp1252"]

        for enc in encodings_to_try:
            try:
                with open(file_path, "r", encoding=enc) as file:
                    content = file.read()
                    return content.strip()
            except UnicodeDecodeError:
                if enc == encodings_to_try[-1]:
                    logger.error(f"Could not decode file with any encoding: {file_path}")
                    raise
                continue
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {str(e)}")
                raise
        return " "


class CVReader(FileReader):
    """Specialized reader for CV files with additional metadata extraction."""

    @staticmethod
    def read_cv(file_path: str) -> Dict[str, Any]:
        """
        Read CV file and extract basic information.

        Args:
            file_path: Path to CV file

        Returns:
            Dictionary with CV content and metadata
        """
        content = FileReader.read_file(file_path)
        file_info = FileReader.get_file_info(file_path)

        return {
            "content": content,
            "file_info": file_info,
            "word_count": len(content.split()),
            "char_count": len(content),
            "has_email": bool(
                re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", content)
            ),
            "has_phone": bool(re.search(r"\b\d{3}[-.]?\d{3}[-.]?\d{4}\b", content)),
            "has_linkedin": "linkedin" in content.lower(),
            "has_github": "github" in content.lower(),
        }


class JobDescriptionReader(FileReader):
    """Specialized reader for job description files."""

    @staticmethod
    def read_job(file_path: str) -> Dict[str, Any]:
        """
        Read job description file and extract basic information.

        Args:
            file_path: Path to job description file

        Returns:
            Dictionary with job content and metadata
        """
        content = FileReader.read_file(file_path)
        file_info = FileReader.get_file_info(file_path)

        # Extract basic job info
        company_match = re.search(r"company[:\s]+([^\n]+)", content, re.IGNORECASE)
        position_match = re.search(r"position[:\s]+([^\n]+)", content, re.IGNORECASE)
        location_match = re.search(r"location[:\s]+([^\n]+)", content, re.IGNORECASE)

        return {
            "content": content,
            "file_info": file_info,
            "word_count": len(content.split()),
            "char_count": len(content),
            "company": company_match.group(1).strip() if company_match else None,
            "position": position_match.group(1).strip() if position_match else None,
            "location": location_match.group(1).strip() if location_match else None,
            "has_requirements": "requirements" in content.lower()
            or "qualifications" in content.lower(),
            "has_responsibilities": "responsibilities" in content.lower()
            or "duties" in content.lower(),
        }


def read_cv(file_path: str) -> str:
    """
    Convenience function to read CV file.

    Args:
        file_path: Path to CV file

    Returns:
        CV content as string
    """
    return FileReader.read_file(file_path)


def read_job_description(file_path: str) -> str:
    """
    Convenience function to read job description file.

    Args:
        file_path: Path to job description file

    Returns:
        Job description content as string
    """
    return FileReader.read_file(file_path)


# Example usage
if __name__ == "__main__":
    # Configure logging
    logging.basicConfig(level=logging.INFO)

    # Example: Read a CV
    try:
        cv_data = CVReader.read_cv("data/my_cv.pdf")
        print(f"CV loaded: {cv_data['word_count']} words")
        print(f"Has email: {cv_data['has_email']}")
        print(f"Has LinkedIn: {cv_data['has_linkedin']}")
    except Exception as e:
        print(f"Error: {e}")
